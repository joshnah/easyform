"""
DOCX form filling functionality.

This module handles filling of DOCX forms, including text replacement,
checkbox updates, and font management.
"""

import os
import logging
from typing import List, Optional, Literal
from docx import Document
from .font_manager import get_available_font, get_fonts_cache_dir
from .checkbox_processor import (
    detect_checkbox_entries,
    process_checkbox_entries,
    update_checkbox_in_paragraph,
)
from .fill_processor import detect_fill_entries, process_fill_entries, FillEntry
from .checkbox_processor import CheckboxEntry


def _apply_fills_and_checkboxes(
    # Note: The Document instance is still accessible via the paragraph references in `locations`,
    # so we don't need to pass the full doc object here.
    lines: List[str],
    locations,
    entries: List[FillEntry],
    checkbox_entries: List[CheckboxEntry],
):
    """Internal helper that mutates the underlying `python-docx` objects referenced in *locations* in-place using pre-computed entries."""
    font_cache = {}
    cache_dir = get_fonts_cache_dir()
    # Apply filled_lines back into document
    for entry in entries:
        group = entry.lines.split("\n")
        filled = entry.filled_lines.split("\n")
        n = len(group)
        for i in range(len(lines) - n + 1):
            if lines[i : i + n] == group:
                for j, loc in enumerate(locations[i : i + n]):
                    if loc[0] == "para":
                        para = loc[1]
                        original_font_info = loc[2]
                    else:
                        _, cell, para, original_font_info = loc
                    para.text = filled[j]

                    # Restore original font where possible
                    if original_font_info:
                        original_font_name = original_font_info["name"]
                        if original_font_name not in font_cache:
                            font_name, font_file_path = get_available_font(
                                original_font_name, cache_dir
                            )
                            font_cache[original_font_name] = (font_name, font_file_path)
                        else:
                            font_name, font_file_path = font_cache[original_font_name]
                        for run in para.runs:
                            if font_name != "helv":
                                run.font.name = font_name
                            if original_font_info["size"]:
                                run.font.size = original_font_info["size"]
                            if original_font_info["bold"] is not None:
                                run.font.bold = original_font_info["bold"]
                            if original_font_info["italic"] is not None:
                                run.font.italic = original_font_info["italic"]
                            if original_font_info["underline"] is not None:
                                run.font.underline = original_font_info["underline"]
                break

    # Apply checkbox changes
    for checkbox_entry in checkbox_entries:
        if checkbox_entry.checked_indices is None:
            continue
        context_lines = checkbox_entry.lines.split("\n")
        for i in range(len(lines) - len(context_lines) + 1):
            if lines[i : i + len(context_lines)] == context_lines:
                for checkbox_idx, (rel_line_idx, char_idx) in enumerate(
                    checkbox_entry.checkbox_positions
                ):
                    doc_line_idx = i + rel_line_idx
                    should_check = checkbox_idx in (
                        checkbox_entry.checked_indices or []
                    )
                    location = locations[doc_line_idx]
                    if location[0] == "para":
                        para = location[1]
                        update_checkbox_in_paragraph(para, char_idx, should_check)
                    else:
                        _, _cell, para, _ = location
                        update_checkbox_in_paragraph(para, char_idx, should_check)
                break


def fill_docx_with_entries(
    fill_entries: List[FillEntry],
    checkbox_entries: List[CheckboxEntry],
    form_path: str,
    output_path: Optional[str] = None,
) -> str:
    """Fill a DOCX form by performing a simple find-and-replace for every line in each entry.

    The replacement strategy follows the user's specification:

    1. Split the original placeholder text (``entry.lines``) by newlines to obtain individual *groups*.
    2. Split the replacement text (``entry.filled_lines``) by newlines to obtain corresponding *filled* lines.
    3. Iterate through all paragraphs in the document (including those inside tables) and replace every occurrence
       of ``groups[i]`` with ``filled[i]``.

    This deliberately ignores fonts, runs, and check-box handling in favour of the straightforward Ctrl-F style
    replacement requested by the user.
    """

    doc = Document(form_path)

    # Build a flat list of all paragraph objects (including those in tables)
    paragraphs: List["docx.text.paragraph.Paragraph"] = list(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                paragraphs.extend(cell.paragraphs)

    # Perform replacements
    for entry in fill_entries:
        groups = entry.lines.split("\n")
        filled = entry.filled_lines.split("\n")

        # Ensure we have the same number of replacement lines; if not, pad with empty strings
        if len(filled) < len(groups):
            filled.extend(["" for _ in range(len(groups) - len(filled))])

        for i, group_line in enumerate(groups):
            replacement = filled[i]
            if not group_line:
                # Skip empty search strings to avoid replacing every position
                continue
            for para in paragraphs:
                if group_line in para.text:
                    para.text = para.text.replace(group_line, replacement)

    # Save the modified document
    output_path = output_path or form_path.replace(".docx", "_filled.docx")
    os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)
    doc.save(output_path)
    return output_path


# Rewrite original fill_docx to preserve legacy behaviour but delegate to new implementation


def fill_docx(
    keys: List[str],
    form_path: str,
    context_dir: str,
    output_path: Optional[str],
    placeholder_pattern,
    provider: Literal["openai", "groq", "anythingllm"]
):
    """Legacy DOCX fill. Detects entries and checkboxes and delegates to fill_docx_with_entries."""
    doc = Document(form_path)

    # Extract all lines and track locations with original font information
    lines: List[str] = []
    locations = (
        []
    )  # tuples for replacing: ('para', paragraph, original_font_info) or ('cell', cell, paragraph, original_font_info)

    for para in doc.paragraphs:
        lines.append(para.text)
        # Extract font info from the first run with font information
        original_font_info = None
        for run in para.runs:
            if run.font.name:
                original_font_info = {
                    "name": run.font.name,
                    "size": run.font.size,
                    "bold": run.font.bold,
                    "italic": run.font.italic,
                    "underline": run.font.underline,
                }
                break
        locations.append(("para", para, original_font_info))

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    lines.append(para.text)
                    # Extract font info from the first run with font information
                    original_font_info = None
                    for run in para.runs:
                        if run.font.name:
                            original_font_info = {
                                "name": run.font.name,
                                "size": run.font.size,
                                "bold": run.font.bold,
                                "italic": run.font.italic,
                                "underline": run.font.underline,
                            }
                            break
                    locations.append(("cell", cell, para, original_font_info))

    # Detect fill entries
    entries = detect_fill_entries(lines, keys, placeholder_pattern, provider)
    # Process each entry: infer keys, fill missing context, generate filled_lines
    entries = process_fill_entries(entries, context_dir, placeholder_pattern, provider)

    # # Detect and process checkbox entries
    # checkbox_entries = detect_checkbox_entries(lines, keys)
    # checkbox_entries = process_checkbox_entries(
    #     checkbox_entries, context_dir, keys, provider
    # )

    return fill_docx_with_entries(entries, [], form_path, output_path)
