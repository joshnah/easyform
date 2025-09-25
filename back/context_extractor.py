import os
import glob
import json
import logging
from docx import Document
import pdfplumber
from pdf2image import convert_from_path
from PIL import Image
from .prompts import EXTRACTION_PROMPT_TEMPLATE
from .llm_client import query_gpt
import pytesseract
import cv2
import numpy as np
import concurrent.futures
import threading
import tempfile
from typing import Literal

# Optional PyMuPDF import for embedded image extraction
try:
    import fitz  # type: ignore
except ImportError:
    fitz = None  # type: ignore
    logging.warning(
        "PyMuPDF library not installed, embedded PDF image extraction disabled."
    )


# Add Docling import (optional fallback)
try:
    from docling.document_converter import DocumentConverter
except ImportError:
    DocumentConverter = None  # type: ignore
    logging.warning("Docling library not installed, using fallback extraction methods.")

# Thread-local storage so every worker thread gets its own converter instance.
_thread_local = threading.local()

# Protect first-time instantiation so we don't spawn many heavyweight pipelines at once.
_converter_init_lock = threading.Lock()


def _get_docling_converter():
    """Return a thread-local DocumentConverter (or None if unavailable)."""
    if DocumentConverter is None:
        return None
    if getattr(_thread_local, "converter", None) is None:
        with _converter_init_lock:
            # Double-check to avoid races inside the lock
            if getattr(_thread_local, "converter", None) is None:
                try:
                    _thread_local.converter = DocumentConverter()
                except Exception as e:
                    logging.error(
                        f"Unable to initialise Docling converter in thread: {e}"
                    )
                    _thread_local.converter = None
    return _thread_local.converter


def scan_context_dir(dir_path):
    """Recursively collect docx, pdf, and image files from the directory."""
    patterns = ["*.docx", "*.pdf", "*.png", "*.jpg", "*.jpeg", "*.tiff"]
    files = []
    for pat in patterns:
        files.extend(glob.glob(os.path.join(dir_path, "**", pat), recursive=True))
    return files


def _extract_with_docling(path):
    """Attempt to extract text using Docling. Returns text or None on failure/unavailability."""
    converter = _get_docling_converter()
    if converter is None:
        return None
    try:
        # Each thread has its own converter, so we attempt conversion without a global lock.
        result = converter.convert(path)
        if result and hasattr(result, "document"):
            doc = result.document
            # Prefer markdown export, fall back to raw text if available
            if hasattr(doc, "export_to_markdown"):
                return doc.export_to_markdown()
            elif hasattr(doc, "export_to_text"):
                return doc.export_to_text()
        return None
    except Exception as e:
        logging.error(f"Docling failed to extract {path}: {e}")
        return None


def _clean_extracted_text(text: str) -> str:
    """Remove placeholder comments and newlines, then strip.

    Docling sometimes yields only HTML comments such as ``<!-- image -->`` when
    it detects pages that contain nothing but images.  After removing these
    placeholders (and line-breaks) the remaining string might be empty; in that
    case we should treat the extraction as if it contained no text so that the
    pipeline can fall back to image-based OCR or other strategies.
    """
    if not text:
        return ""
    return text.replace("<!-- image -->", "").replace("\n", "").strip()


def extract_docx(path):
    """Extract text from DOCX file, preferring Docling when available."""
    # Try Docling first for unified extraction
    docling_text = _extract_with_docling(path)
    if docling_text and _clean_extracted_text(docling_text):
        return docling_text
    # Fallback to legacy extraction
    try:
        doc = Document(path)
        texts = [p.text for p in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    texts.append(cell.text)
        return "\n".join(texts)
    except Exception as e:
        logging.error(f"Error extracting DOCX {path}: {e}")
        return ""


def preprocess_for_ocr(pil_img):
    """Preprocess a PIL Image for better OCR accuracy: denoise, thresholding, and scaling."""
    # Convert PIL image to OpenCV format (BGR)
    img = cv2.cvtColor(np.array(pil_img), cv2.COLOR_RGB2BGR)
    # Convert to grayscale
    gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
    # Apply median blur to reduce noise
    blur = cv2.medianBlur(gray, 3)
    # Adaptive thresholding to enhance text regions
    thresh = cv2.adaptiveThreshold(
        blur, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, cv2.THRESH_BINARY, 11, 2
    )
    # Convert back to PIL image
    return Image.fromarray(thresh)


def _extract_images_from_pdf(path: str):
    """Extract embedded images from a PDF and save them to temporary PNG files.

    Returns
    -------
    list[str]
        Paths to the temporary image files. **Caller is responsible for deleting.**
    """
    image_paths = []
    if fitz is None:
        # PyMuPDF not available
        return image_paths

    try:
        doc = fitz.open(path)
        for page_index in range(len(doc)):
            page = doc[page_index]
            for img_index, img_info in enumerate(page.get_images(full=True)):
                xref = img_info[0]
                try:
                    pix = fitz.Pixmap(doc, xref)
                    # Convert CMYK/Gray/other to RGB
                    if pix.alpha or pix.n > 4:
                        pix = fitz.Pixmap(fitz.csRGB, pix)
                    fd, tmp_path = tempfile.mkstemp(suffix=".png")
                    os.close(fd)
                    pix.save(tmp_path)
                    image_paths.append(tmp_path)
                except Exception as e:
                    logging.debug(
                        f"Failed to extract image {img_index} on page {page_index}: {e}"
                    )
        doc.close()
    except Exception as e:
        logging.error(f"Error extracting images from PDF {path}: {e}")

    return image_paths


def _extract_docling_from_pdf_images(pdf_path: str) -> str:
    """Attempt OCR with Docling on all embedded images of the PDF.

    This is our second-chance strategy when direct Docling conversion of the PDF
    yields no usable text. It keeps the pipeline within Docling (avoiding our
    legacy OCR) while working around PDFs that are essentially just image
    containers.
    """
    texts = []
    image_paths = _extract_images_from_pdf(pdf_path)

    for img_path in image_paths:
        try:
            text = _extract_with_docling(img_path) or ""
            cleaned = _clean_extracted_text(text)
            if cleaned:
                texts.append(cleaned)
        except Exception as e:
            logging.debug(f"Docling failed on image {img_path}: {e}")
        finally:
            # Always remove temp file
            try:
                os.remove(img_path)
            except Exception:
                pass

    return "\n".join(texts).strip()


def extract_pdf(path):
    """Extract text from PDF, preferring Docling when available."""
    is_image_only = False
    # Try Docling first
    docling_text = _extract_with_docling(path)
    if docling_text and _clean_extracted_text(docling_text):
        return docling_text, is_image_only

    logging.warning("Docling extraction failed or returned no usable text.")

    # Second attempt: use Docling on each embedded image extracted from the PDF
    img_docling_text = _extract_docling_from_pdf_images(path)
    if img_docling_text and _clean_extracted_text(img_docling_text):
        is_image_only = True
        return img_docling_text, is_image_only

    # Fallback to legacy extraction
    logging.info(f"Falling back to pdfplumber for text extraction: {path}")
    try:
        # First try pdfplumber for text extraction
        with pdfplumber.open(path) as pdf:
            plumber_texts = []
            for page in pdf.pages:
                text = page.extract_text() or ""
                plumber_texts.append(text)
            plumber_result = "\n".join(plumber_texts).strip()
            if plumber_result and _clean_extracted_text(plumber_result):
                return plumber_result, is_image_only
        # Fallback: convert PDF to images and use OCR
        logging.info(
            f"No text found with pdfplumber, converting PDF to images for OCR: {path}"
        )
        is_image_only = True
        images = convert_from_path(path, dpi=200)
        texts = []
        for i, image in enumerate(images):
            logging.debug(f"Processing page {i+1} of {len(images)} from PDF")
            text = extract_image_from_pil(image)
            texts.append(text)
        return "\n".join(texts), is_image_only
    except Exception as e:
        logging.error(f"Error extracting PDF {path}: {e}")
        return "", is_image_only


def extract_image(path):
    """Perform OCR on an image file, preferring Docling when available."""
    # Try Docling first
    docling_text = _extract_with_docling(path)
    if docling_text and _clean_extracted_text(docling_text):
        return docling_text
    # Fallback to legacy OCR pipeline
    try:
        # Load and upscale image
        img = Image.open(path).convert("RGB")
        scale_factor = 3
        new_size = (img.width * scale_factor, img.height * scale_factor)
        img_upscaled = img.resize(new_size, resample=Image.LANCZOS)
        # Convert to OpenCV grayscale
        img_np = cv2.cvtColor(np.array(img_upscaled), cv2.COLOR_RGB2GRAY)
        # Denoise with bilateral filter to preserve edges
        denoised = cv2.bilateralFilter(img_np, 9, 75, 75)
        # Apply CLAHE for contrast enhancement
        clahe = cv2.createCLAHE(clipLimit=2.0, tileGridSize=(8, 8))
        equalized = clahe.apply(denoised)
        # Sharpen image
        kernel = np.array([[0, -1, 0], [-1, 5, -1], [0, -1, 0]])
        sharpened = cv2.filter2D(equalized, -1, kernel)
        proc_img = Image.fromarray(sharpened)
        # Perform OCR with Tesseract
        ocr_config = "--oem 3 --psm 6"
        text = pytesseract.image_to_string(proc_img, config=ocr_config)
        # Fallback: threshold-based preprocessing if no text
        if not text.strip():
            thresh_img = preprocess_for_ocr(img_upscaled)
            text = pytesseract.image_to_string(thresh_img, config=ocr_config)
        # Fallback: single-line mode
        if not text.strip():
            text = pytesseract.image_to_string(proc_img, config="--oem 3 --psm 7")
        return text
    except Exception as e:
        logging.error(f"Error OCR image {path}: {e}")
        return ""


def extract_image_from_pil(img):
    """Perform OCR on a PIL Image with enhanced DPI scaling and specialized preprocessing."""
    try:
        # Convert to RGB if not already
        img = img.convert("RGB")
        scale_factor = 3
        new_size = (img.width * scale_factor, img.height * scale_factor)
        img_upscaled = img.resize(new_size, resample=Image.LANCZOS)
        # Convert to OpenCV grayscale
        img_np = cv2.cvtColor(np.array(img_upscaled), cv2.COLOR_RGB2GRAY)
        # Denoise with bilateral filter to preserve edges
        denoised = cv2.bilateralFilter(img_np, 9, 75, 75)
        # Apply CLAHE for contrast enhancement
        clahe = cv2.createCLAHE(clipLimit=2.0, tileGridSize=(8, 8))
        equalized = clahe.apply(denoised)
        # Sharpen image
        kernel = np.array([[0, -1, 0], [-1, 5, -1], [0, -1, 0]])
        sharpened = cv2.filter2D(equalized, -1, kernel)
        proc_img = Image.fromarray(sharpened)
        # Perform OCR with Tesseract
        ocr_config = "--oem 3 --psm 6"
        text = pytesseract.image_to_string(proc_img, config=ocr_config)
        # Fallback: threshold-based preprocessing if no text
        if not text.strip():
            thresh_img = preprocess_for_ocr(img_upscaled)
            text = pytesseract.image_to_string(thresh_img, config=ocr_config)
        # Fallback: single-line mode
        if not text.strip():
            text = pytesseract.image_to_string(proc_img, config="--oem 3 --psm 7")
        return text
    except Exception as e:
        logging.error(f"Error OCR PIL image: {e}")
        return ""


def aggregate_text(paths):
    """Aggregate extracted text from a list of file paths."""
    corpus = []
    for path in paths:
        ext = os.path.splitext(path)[1].lower()
        logging.info(f"Extracting text from {path}")
        if ext == ".docx":
            corpus.append(extract_docx(path))
        elif ext == ".pdf":
            # extract_pdf returns a tuple (text, is_image_only). We only need the text for the corpus.
            pdf_result = extract_pdf(path)
            # Be defensive: handle both tuple and plain string returns.
            if isinstance(pdf_result, tuple):
                text, _ = pdf_result
            else:
                text = pdf_result
            corpus.append(text)
        elif ext in [".png", ".jpg", ".jpeg", ".tiff"]:
            corpus.append(extract_image(path))
    return "\n".join(corpus)


def extract_personal_info(raw_text, provider: Literal["openai", "groq", "anythingllm"], lang='en'):
    """Use the LLM to extract personal info JSON from raw text."""
    prompt = EXTRACTION_PROMPT_TEMPLATE.format(content=raw_text)
    response = query_gpt(prompt, provider=provider)
    
    # Log the raw response for debugging
    logging.debug(f"Raw GPT response: {repr(response)}")

    if not response or not response.strip():
        logging.error("Empty response from GPT")
        raise ValueError("Empty response from GPT")

    try:
        return json.loads(response)
    except json.JSONDecodeError as e:
        logging.debug(f"Direct JSON parsing failed: {e}")
        logging.debug(f"Attempting to extract JSON from wrapped response")

        # Try to extract JSON from the response if it's wrapped in markdown or other text
        import re

        # First try to extract from markdown code blocks
        markdown_match = re.search(
            r"```(?:json)?\s*(\{.*?\})\s*```", response, re.DOTALL
        )
        if markdown_match:
            try:
                extracted_json = markdown_match.group(1)
                logging.debug(
                    f"Extracted JSON from markdown: {extracted_json[:100]}..."
                )
                return json.loads(extracted_json)
            except json.JSONDecodeError:
                logging.debug("Failed to parse extracted JSON from markdown")

        # Fallback to finding any JSON object in the response
        json_match = re.search(r"\{.*\}", response, re.DOTALL)
        if json_match:
            try:
                extracted_json = json_match.group()
                logging.debug(
                    f"Extracted JSON from general search: {extracted_json[:100]}..."
                )
                return json.loads(extracted_json)
            except json.JSONDecodeError:
                logging.debug("Failed to parse extracted JSON from general search")

        # If all else fails, return a default structure
        logging.error(f"Failed to parse JSON from GPT response: {e}")
        logging.error(f"Response content: {repr(response)}")
        logging.warning(
            "Returning empty personal info structure due to parsing failure"
        )
        return {
            "full_name": "",
            "first_name": "",
            "middle_names": "",
            "last_name": "",
            "birth_day": "",
            "birth_month": "",
            "birth_year": "",
            "date_of_birth (MM-DD-YYYY)": "",
            "date_of_birth (DD-MM-YYYY)": "",
            "date_of_birth (MM/DD/YYYY)": "",
            "date_of_birth (DD/MM/YYYY)": "",
            "date_of_birth (YYYY/MM/DD)": "",
            "date_of_birth (YYYY-MM-DD)": "",
            "phone_number": "",
            "email": "",
            "address": "",
        }


def get_source_type(file_path):
    """Determine the source type based on file extension."""
    ext = os.path.splitext(file_path)[1].lower()
    if ext == ".docx":
        return "text"
    elif ext == ".pdf":
        return "pdf"
    elif ext in [".png", ".jpg", ".jpeg", ".tiff"]:
        return "image"
    else:
        return "unknown"


def extract_from_individual_files(file_paths, provider: Literal["openai", "groq", "anythingllm"], lang='en'):
    """Extract personal info from each file individually, tracking source types, in parallel."""
    extractions = []

    def process_file(file_path):
        source_type = get_source_type(file_path)
        logging.info(f"Processing {source_type} file: {file_path}")
        ext = os.path.splitext(file_path)[1].lower()
        if ext == ".docx":
            text_content = extract_docx(file_path)
        elif ext == ".pdf":
            text_content, is_image_only = extract_pdf(file_path)
            if is_image_only:
                source_type = "image"
        elif ext in [".png", ".jpg", ".jpeg", ".tiff"]:
            text_content = extract_image(file_path)
        else:
            logging.warning(f"Unsupported file type: {file_path}")
            return None
        if not text_content.strip():
            logging.warning(f"No text extracted from {file_path}")
            return None
        try:
            extracted_data = extract_personal_info(text_content, provider, lang)
            logging.debug(f"Extracted from {file_path}: {extracted_data}")
            return {
                "source_file": file_path,
                "source_type": source_type,
                "extracted_data": extracted_data,
            }
        except Exception as e:
            logging.error(f"Failed to extract from {file_path}: {e}")
            return None

    with concurrent.futures.ThreadPoolExecutor() as executor:
        futures = [executor.submit(process_file, file_path) for file_path in file_paths]
        for future in concurrent.futures.as_completed(futures):
            result = future.result()
            if result is not None:
                extractions.append(result)

    return extractions


def _resolve_most_frequent(values, source_type, field):
    """Helper to resolve the most frequent value from a list of dicts with 'value' and 'source'."""
    if len(values) == 1:
        final_value = values[0]["value"]
        resolution = f"{source_type}_single from {values[0]['source']}"
        conflict_log = None
    else:
        value_counts = {}
        lower_values_to_original = {}
        for item in values:
            val = item["value"]
            if val.lower() not in value_counts:
                value_counts[val.lower()] = []
                lower_values_to_original[val.lower()] = val
            value_counts[val.lower()].append(item["source"])
        most_frequent = max(value_counts.keys(), key=lambda x: len(value_counts[x]))
        final_value = lower_values_to_original[most_frequent]
        resolution = (
            f"{source_type}_frequent '{final_value}' from {value_counts[most_frequent]}"
        )
        conflict_log = None
        if len(value_counts) > 1:
            conflict_log = f"Field '{field}': Multiple {source_type} values found, chose most frequent '{most_frequent}'"
    return final_value, resolution, conflict_log


def resolve_conflicts(extractions):
    """Resolve conflicts between multiple extractions using priority rules."""
    if not extractions:
        return {
            "full_name": "",
            "first_name": "",
            "middle_names": "",
            "last_name": "",
            "birth_day": "",
            "birth_month": "",
            "birth_year": "",
            "date_of_birth (MM-DD-YYYY)": "",
            "date_of_birth (DD-MM-YYYY)": "",
            "date_of_birth (MM/DD/YYYY)": "",
            "date_of_birth (DD/MM/YYYY)": "",
            "date_of_birth (YYYY/MM/DD)": "",
            "date_of_birth (YYYY-MM-DD)": "",
            "phone_number": "",
            "email": "",
            "address": "",
        }

    # Group extractions by field
    field_values = {}
    field_names = [
        "full_name",
        "first_name",
        "middle_names",
        "last_name",
        "birth_day",
        "birth_month",
        "birth_year",
        "date_of_birth (MM-DD-YYYY)",
        "date_of_birth (DD-MM-YYYY)",
        "date_of_birth (MM/DD/YYYY)",
        "date_of_birth (DD/MM/YYYY)",
        "date_of_birth (YYYY/MM/DD)",
        "date_of_birth (YYYY-MM-DD)",
        "phone_number",
        "email",
        "address",
    ]

    for field in field_names:
        field_values[field] = {"text": [], "pdf": [], "image": []}

    # Collect all values for each field by source type
    for extraction in extractions:
        source_type = extraction["source_type"]
        data = extraction["extracted_data"]
        source_file = extraction["source_file"]

        for field in field_names:
            value = data.get(field, "").strip()
            if value:  # Only collect non-empty values
                field_values[field][source_type].append(
                    {"value": value, "source": source_file}
                )

    # Resolve conflicts for each field
    final_result = {}
    conflicts_log = []

    for field in field_names:
        values = field_values[field]
        text_values = values["text"]
        pdf_values = values["pdf"]
        image_values = values["image"]

        final_value = ""
        resolution_info = {
            "field": field,
            "text_count": len(text_values),
            "pdf_count": len(pdf_values),
            "image_count": len(image_values),
            "resolution": "empty",
        }

        if text_values:
            final_value, resolution, conflict_log = _resolve_most_frequent(
                text_values, "text", field
            )
            resolution_info["resolution"] = resolution
            if conflict_log:
                conflicts_log.append(conflict_log)
        elif pdf_values:
            final_value, resolution, conflict_log = _resolve_most_frequent(
                pdf_values, "pdf", field
            )
            resolution_info["resolution"] = resolution
            if conflict_log:
                conflicts_log.append(conflict_log)
        elif image_values:
            final_value, resolution, conflict_log = _resolve_most_frequent(
                image_values, "image", field
            )
            resolution_info["resolution"] = resolution
            if conflict_log:
                conflicts_log.append(conflict_log)
        # If no values found, use empty string
        if not final_value:
            final_value = ""
            resolution_info["resolution"] = "empty"
        final_result[field] = final_value
        logging.debug(f"Field '{field}' resolution: {resolution_info}")

    # Log conflict resolutions
    if conflicts_log:
        logging.info("Conflict resolutions applied:")
        for conflict in conflicts_log:
            logging.info(f"  - {conflict}")
    else:
        logging.info("No conflicts found between sources")

    return final_result


def extract_context(dir_path, provider: Literal["openai", "groq", "anythingllm"], lang='en'):
    """Full context extraction pipeline with conflict resolution."""
    files = scan_context_dir(dir_path)

    if not files:
        logging.warning(f"No supported files found in {dir_path}")
        return {
            "full_name": "",
            "first_name": "",
            "middle_names": "",
            "last_name": "",
            "birth_day": "",
            "birth_month": "",
            "birth_year": "",
            "date_of_birth (MM-DD-YYYY)": "",
            "date_of_birth (DD-MM-YYYY)": "",
            "date_of_birth (MM/DD/YYYY)": "",
            "date_of_birth (DD/MM/YYYY)": "",
            "date_of_birth (YYYY/MM/DD)": "",
            "date_of_birth (YYYY-MM-DD)": "",
            "phone_number": "",
            "email": "",
            "address": "",
        }

    logging.info(f"Found {len(files)} files to process")

    # Extract from each file individually
    extractions = extract_from_individual_files(files, provider, lang)
    
    if not extractions:
        logging.warning("No successful extractions from any files")
        return {
            "full_name": "",
            "first_name": "",
            "middle_names": "",
            "last_name": "",
            "birth_day": "",
            "birth_month": "",
            "birth_year": "",
            "date_of_birth (MM-DD-YYYY)": "",
            "date_of_birth (DD-MM-YYYY)": "",
            "date_of_birth (MM/DD/YYYY)": "",
            "date_of_birth (DD/MM/YYYY)": "",
            "date_of_birth (YYYY/MM/DD)": "",
            "date_of_birth (YYYY-MM-DD)": "",
            "phone_number": "",
            "email": "",
            "address": "",
        }

    logging.info(f"Successfully extracted from {len(extractions)} files")

    # Resolve conflicts and merge results
    final_result = resolve_conflicts(extractions)

    # Log final summary
    non_empty_fields = {k: v for k, v in final_result.items() if v.strip()}
    logging.info(
        f"Final extraction completed with {len(non_empty_fields)} populated fields: {list(non_empty_fields.keys())}"
    )

    return final_result
