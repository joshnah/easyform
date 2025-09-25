"""
Text extraction utilities for form filling.

This module handles extraction of text content from various file formats
including DOCX and PDF files for pattern analysis and processing.
"""

import os
import logging
from docx import Document
from pypdf import PdfReader
import fitz


def extract_form_text(form_path: str) -> str:
    """Extract all text from a form file (DOCX or PDF) for pattern analysis."""
    ext = os.path.splitext(form_path)[1].lower()
    
    if ext == '.docx':
        doc = Document(form_path)
        lines = []
        
        # Extract from paragraphs
        for para in doc.paragraphs:
            lines.append(para.text)
        
        # Extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        lines.append(para.text)
        
        return '\n'.join(lines)
    
    elif ext == '.pdf':
        try:
            # Try with PyMuPDF first
            doc = fitz.open(form_path)
            text_parts = []
            for page in doc:
                text_parts.append(page.get_text())
            doc.close()
            return '\n'.join(text_parts)
        except Exception as e:
            logging.warning(f"PyMuPDF failed: {e}, trying pypdf")
            # Fallback to pypdf
            try:
                reader = PdfReader(form_path)
                text_parts = []
                for page in reader.pages:
                    text_parts.append(page.extract_text())
                return '\n'.join(text_parts)
            except Exception as e2:
                logging.error(f"Both PDF readers failed: {e2}")
                return ""
    
    else:
        raise ValueError(f"Unsupported form format: {ext}") 